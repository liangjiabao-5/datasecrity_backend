import os
import sys
import re
import subprocess
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT, CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import ns
from docx.shared import Cm
from copy import deepcopy

# Add src to path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.indexer import Indexer
from src.searcher import Searcher

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF')
    rPr.append(u); rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def insert_attachments(doc, paragraph, token_path):
    if not os.path.exists(token_path):
        print(f"附件不存在: {token_path}")
        paragraph.add_run(f"[附件缺失: {token_path}]")
        return
    if 'o' not in ns.nsmap:
        ns.nsmap['o'] = 'urn:schemas-microsoft-com:office:office'
    ext = os.path.splitext(token_path)[1].lower()
    if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
        try:
            # 直接在当前段落之后创建图片段落
            pic_par = doc.add_paragraph()
            pic_run = pic_par.add_run()
            pic = pic_run.add_picture(token_path)
            
            # 在图片段落之后创建标题段落
            cap_par = doc.add_paragraph()
            cap_par.add_run(f"（图片：{os.path.basename(token_path)}）")
            
            # 获取两个段落的XML元素
            pic_el = pic_par._p
            cap_el = cap_par._p
            
            # 从文档中移除这两个段落
            doc.element.body.remove(pic_el)
            doc.element.body.remove(cap_el)
            
            # 将图片段落插入到当前段落之后
            paragraph._p.addnext(pic_el)
            
            # 将标题段落插入到图片段落之后
            pic_el.addnext(cap_el)
        except Exception as e:
            print(f"图片嵌入异常: {e}")
            paragraph.add_run(f"[图片嵌入异常: {os.path.basename(token_path)}]")
    elif ext == '.docx':
        try:
            sub = Document(token_path)
            # 深度复制表格的底层 XML 以保留原始格式（合并、边框、字体、背景等）
            anchor = paragraph._p
            inserted = False
            for el in list(sub.element.body):
                if el.tag == qn('w:tbl'):
                    new_el = deepcopy(el)
                    anchor.addnext(new_el)
                    anchor = new_el
                    inserted = True
            if inserted:
                cap_p = doc.paragraphs[-1]
                cap_p.add_run(f"（表格：{os.path.basename(token_path)}）")
            else:
                paragraph.add_run(f"[未发现表格元素：{os.path.basename(token_path)}]")
        except Exception:
            paragraph.add_run(f"[表格嵌入异常: {os.path.basename(token_path)}]")
    elif ext in ['.vsd', '.vsdx', '.vss', '.vssx', '.vdx', '.vsb', '.bin']:
        try:
            pkg = doc.part.package
            # 为Visio文件创建正确的partname，根据文件扩展名
            visio_ext = os.path.splitext(token_path)[1].lower()
            partname = pkg.next_partname(f"/word/embeddings/visio%d{visio_ext}")
            with open(token_path, "rb") as f:
                blob = f.read()
            
            # 确定正确的内容类型
            if visio_ext == '.vsdx':
                content_type = 'application/vnd.ms-visio.drawing'
            else:
                content_type = 'application/vnd.ms-visio.application'
            
            ole_part = Part(PackURI(partname), content_type, blob, pkg)
            r_id = doc.part.relate_to(ole_part, RT.OLE_OBJECT)
            
            print(f"嵌入 Visio OLE 对象: {token_path} -> {partname}, rId={r_id}")
            
            # 构造标准的 OLE 对象 XML 结构
            obj = OxmlElement("w:object")
            
            # 使用 v:shape 作为容器，添加必要的命名空间
            if 'v' not in ns.nsmap:
                ns.nsmap['v'] = 'urn:schemas-microsoft-com:vml'
            if 'o' not in ns.nsmap:
                ns.nsmap['o'] = 'urn:schemas-microsoft-com:office:office'
            if 'w' not in ns.nsmap:
                ns.nsmap['w'] = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            shape = OxmlElement('v:shape')
            shape_id = f'_x0000_i{r_id[3:]}'
            shape.set('id', shape_id)
            shape.set('type', '#_x0000_t75')
            shape.set('style', 'width:400pt;height:300pt') # 更大的默认大小，便于查看
            shape.set(qn('o:ole'), '')
            
            # 添加 v:imagedata 元素，使用占位符图像
            # 虽然没有实际预览图，但添加这个元素可以避免显示黑框
            imagedata = OxmlElement('v:imagedata')
            imagedata.set(qn('o:relid'), r_id)
            imagedata.set('src', f'visio_placeholder{r_id[3:]}.png')
            shape.append(imagedata)
            
            obj.append(shape)
            
            # 创建更完整的 OLE 对象元素
            ole = OxmlElement("o:OLEObject")
            ole.set(qn("r:id"), r_id)
            # 根据文件类型设置正确的 ProgID
            if visio_ext == '.vsdx':
                ole.set("ProgID", "Visio.Drawing.15")  # 适用于较新的 Visio 版本
            else:
                ole.set("ProgID", "Visio.Drawing.12")  # 适用于旧版本
            ole.set("Type", "Embed")
            ole.set("ShapeID", shape_id)
            ole.set("DrawAspect", "Content")
            ole.set("ObjectID", f"_obj{r_id[3:]}")
            obj.append(ole)
            
            # 创建正确的段落和运行元素
            p_obj = OxmlElement("w:p")
            p_pr = OxmlElement("w:pPr")
            p_obj.append(p_pr)
            r_obj = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            r_obj.append(r_pr)
            r_obj.append(obj)
            p_obj.append(r_obj)
            paragraph._p.addnext(p_obj)
            
            cap_p = doc.paragraphs[-1]
            cap_p.add_run(f"（Visio对象：{os.path.basename(token_path)}）")
        except Exception as e:
            print(f"Visio 嵌入异常: {token_path}, error={e}")
            paragraph.add_run(f"[Visio嵌入异常: {os.path.basename(token_path)}]")
    else:
        paragraph.add_run(f"[附件: {token_path}]")

def export_results_to_docx(results, output_path):
    doc = Document()
    doc.add_heading('检索结果', level=1)
    if not results:
        doc.add_paragraph("无检索结果")
    else:
        res = results[0]
        doc.add_heading(f"结果 1（Score: {res['score']:.4f}）", level=2)
        doc.add_paragraph(f"文件：{res['filename']}")
        doc.add_paragraph(f"标题：{res['title']}")
        content = res['content']
        current_par = doc.add_paragraph()
        pos = 0
        for m in re.finditer(r'\{\{ATTACHMENT:(.*?)\}\}', content):
            before = content[pos:m.start()]
            if before:
                for line in before.splitlines():
                    current_par.add_run(line)
                    current_par = doc.add_paragraph()
            path = m.group(1)
            insert_attachments(doc, current_par, path)
            current_par = doc.add_paragraph()
            pos = m.end()
        tail = content[pos:]
        if tail:
            for line in tail.splitlines():
                current_par.add_run(line)
                current_par = doc.add_paragraph()
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved Word results to: {output_path}")

def main():
    print("=== RAG System Demo ===")
    
    # 1. Indexing
    print("\n[Step 1] Indexing documents...")
    indexer = Indexer()
    # Assuming data is in 'data' folder
    data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        print(f"Created {data_dir}. Please put your files there.")
        return

    indexer.index_directory(data_dir)
    
    # 2. Searching
    print("\n[Step 2] Searching...")
    searcher = Searcher()
    
    while True:
        print("\n" + "-"*30)
        query = input("Enter query (or 'q' to quit): ").strip()
        if query.lower() == 'q':
            break
        
        filter_file = input("Filter by filename (optional, press Enter to skip): ").strip()
        filter_title = input("Filter by title (optional, press Enter to skip): ").strip()
        filter_content = input("Filter by content (optional, press Enter to skip): ").strip()
        
        results = searcher.search(
            query, 
            filename_filter=filter_file if filter_file else None,
            title_filter=filter_title if filter_title else None,
            content_filter=filter_content if filter_content else None
        )
        
        print(f"\nFound {len(results)} results:")
        # 显示搜索结果预览
        for i, result in enumerate(results):
            print(f"\nResult {i+1} (Score: {result['score']:.4f})")
            print(f"Filename: {result['filename']}")
            print(f"Title: {result['title']}")
            print(f"Preview: {result['preview']}")
        
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
        output_path = os.path.join(output_dir, 'search_top1.docx')
        export_results_to_docx(results, output_path)

if __name__ == "__main__":
    main()
